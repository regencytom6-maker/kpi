from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.db.models import Q, Count
from datetime import datetime, timedelta
from bmr.models import BMR, BMRSignature
from workflow.models import BatchPhaseExecution
import csv
import json

@login_required
def comments_report_view(request):
    """Web-based comments report view with role-based filtering"""
    
    # Check if user is admin/staff - they see all comments
    is_admin = request.user.is_staff or request.user.is_superuser or request.user.role == 'admin'
    
    # Collect all comments
    comments_data = []
    
    # 1. BMR Level Comments
    if is_admin:
        # Admin sees all BMRs
        bmrs = BMR.objects.all().select_related('product', 'created_by', 'approved_by')
    else:
        # Operators only see BMRs they created or were involved in
        bmrs = BMR.objects.filter(
            Q(created_by=request.user) | 
            Q(approved_by=request.user)
        ).select_related('product', 'created_by', 'approved_by')
    
    for bmr in bmrs:
        # QA Comments on BMR
        if bmr.qa_comments:
            comments_data.append({
                'bmr_number': bmr.batch_number,
                'product': bmr.product.product_name,
                'comment_type': 'BMR QA Comments',
                'phase': 'BMR Creation',
                'user': bmr.created_by.get_full_name() if bmr.created_by else 'Unknown',
                'user_role': bmr.created_by.role if bmr.created_by else 'Unknown',
                'date': bmr.created_date,
                'comments': bmr.qa_comments,
                'status': bmr.status,
                'bmr_id': bmr.id
            })
        
        # Regulatory Comments on BMR
        if bmr.regulatory_comments:
            comments_data.append({
                'bmr_number': bmr.batch_number,
                'product': bmr.product.product_name,
                'comment_type': 'BMR Regulatory Comments',
                'phase': 'Regulatory Approval',
                'user': bmr.approved_by.get_full_name() if bmr.approved_by else 'Unknown',
                'user_role': bmr.approved_by.role if bmr.approved_by else 'Unknown',
                'date': bmr.approved_date or bmr.modified_date,
                'comments': bmr.regulatory_comments,
                'status': bmr.status,
                'bmr_id': bmr.id
            })
    
    # 2. Phase Level Comments
    if is_admin:
        # Admin sees all phases
        phases = BatchPhaseExecution.objects.all().select_related(
            'bmr', 'bmr__product', 'phase', 'started_by', 'completed_by'
        )
    else:
        # Operators only see phases they were involved in
        phases = BatchPhaseExecution.objects.filter(
            Q(started_by=request.user) | 
            Q(completed_by=request.user) |
            Q(bmr__created_by=request.user)
        ).select_related(
            'bmr', 'bmr__product', 'phase', 'started_by', 'completed_by'
        )
    
    for phase in phases:
        # Operator Comments
        if phase.operator_comments:
            comments_data.append({
                'bmr_number': phase.bmr.batch_number,
                'product': phase.bmr.product.product_name,
                'comment_type': 'Operator Comments',
                'phase': phase.phase.get_phase_name_display(),
                'user': phase.completed_by.get_full_name() if phase.completed_by else 
                        phase.started_by.get_full_name() if phase.started_by else 'Unknown',
                'user_role': (phase.completed_by.role if phase.completed_by else 
                            phase.started_by.role if phase.started_by else 'Unknown'),
                'date': phase.completed_date or phase.started_date or phase.created_date,
                'comments': phase.operator_comments,
                'status': phase.status,
                'bmr_id': phase.bmr.id,
                'phase_id': phase.id
            })
        
        # QA Comments on Phase
        if phase.qa_comments:
            comments_data.append({
                'bmr_number': phase.bmr.batch_number,
                'product': phase.bmr.product.product_name,
                'comment_type': 'Phase QA Comments',
                'phase': phase.phase.get_phase_name_display(),
                'user': phase.completed_by.get_full_name() if phase.completed_by else 'Unknown',
                'user_role': phase.completed_by.role if phase.completed_by else 'Unknown',
                'date': phase.completed_date or phase.created_date,
                'comments': phase.qa_comments,
                'status': phase.status,
                'bmr_id': phase.bmr.id,
                'phase_id': phase.id
            })
        
        # Rejection Reasons
        if phase.rejection_reason:
            comments_data.append({
                'bmr_number': phase.bmr.batch_number,
                'product': phase.bmr.product.product_name,
                'comment_type': 'Rejection Reason',
                'phase': phase.phase.get_phase_name_display(),
                'user': phase.completed_by.get_full_name() if phase.completed_by else 'Unknown',
                'user_role': phase.completed_by.role if phase.completed_by else 'Unknown',
                'date': phase.completed_date or phase.created_date,
                'comments': phase.rejection_reason,
                'status': phase.status,
                'bmr_id': phase.bmr.id,
                'phase_id': phase.id
            })
    
    # 3. Signature Comments
    if is_admin:
        # Admin sees all signatures
        signatures = BMRSignature.objects.all().select_related('bmr', 'bmr__product', 'signed_by')
    else:
        # Operators only see signatures they made or on BMRs they created
        signatures = BMRSignature.objects.filter(
            Q(signed_by=request.user) |
            Q(bmr__created_by=request.user)
        ).select_related('bmr', 'bmr__product', 'signed_by')
    
    for signature in signatures:
        if signature.comments:
            comments_data.append({
                'bmr_number': signature.bmr.batch_number,
                'product': signature.bmr.product.product_name,
                'comment_type': 'Electronic Signature',
                'phase': f"Signature - {signature.signed_by_role}",
                'user': signature.signed_by.get_full_name() if signature.signed_by else 'Unknown',
                'user_role': signature.signed_by_role,
                'date': signature.signed_date,
                'comments': signature.comments,
                'status': 'Signed',
                'bmr_id': signature.bmr.id,
                'signature_id': signature.id
            })
    
    # Sort by date (newest first)
    comments_data.sort(key=lambda x: x['date'] or datetime.min, reverse=True)
    
    # Filter by request parameters
    bmr_filter = request.GET.get('bmr')
    comment_type_filter = request.GET.get('type')
    user_role_filter = request.GET.get('role')
    
    if bmr_filter:
        comments_data = [c for c in comments_data if bmr_filter.lower() in c['bmr_number'].lower()]
    
    if comment_type_filter:
        comments_data = [c for c in comments_data if c['comment_type'] == comment_type_filter]
    
    if user_role_filter:
        comments_data = [c for c in comments_data if c['user_role'] == user_role_filter]
    
    # Generate statistics
    stats = {
        'total_comments': len(comments_data),
        'bmr_comments': len([c for c in comments_data if 'BMR' in c['comment_type']]),
        'phase_comments': len([c for c in comments_data if c['comment_type'] in ['Operator Comments', 'Phase QA Comments']]),
        'rejections': len([c for c in comments_data if c['comment_type'] == 'Rejection Reason']),
        'signatures': len([c for c in comments_data if c['comment_type'] == 'Electronic Signature']),
    }
    
    # Get unique values for filters
    all_comment_types = list(set([c['comment_type'] for c in comments_data]))
    all_user_roles = list(set([c['user_role'] for c in comments_data]))
    all_bmrs = list(set([c['bmr_number'] for c in comments_data]))
    
    context = {
        'comments': comments_data[:100],  # Limit to first 100 for performance
        'total_comments': len(comments_data),
        'stats': stats,
        'comment_types': sorted(all_comment_types),
        'user_roles': sorted(all_user_roles),
        'bmrs': sorted(all_bmrs),
        'current_filters': {
            'bmr': bmr_filter,
            'type': comment_type_filter,
            'role': user_role_filter
        },
        'is_admin': is_admin,  # Pass admin status to template
        'user_role': request.user.role  # Pass user role to template
    }
    
    return render(request, 'reports/comments_report.html', context)

@login_required
def export_comments_csv(request):
    """Export comments to CSV format with role-based filtering"""
    
    # Check if user is admin/staff - they see all comments
    is_admin = request.user.is_staff or request.user.is_superuser or request.user.role == 'admin'
    
    # Collect all comments (same logic as above)
    comments_data = []
    
    # BMR Level Comments
    if is_admin:
        bmrs = BMR.objects.all().select_related('product', 'created_by', 'approved_by')
    else:
        bmrs = BMR.objects.filter(
            Q(created_by=request.user) | 
            Q(approved_by=request.user)
        ).select_related('product', 'created_by', 'approved_by')
    
    for bmr in bmrs:
        if bmr.qa_comments:
            comments_data.append({
                'BMR Number': bmr.batch_number,
                'Product': bmr.product.product_name,
                'Comment Type': 'BMR QA Comments',
                'Phase': 'BMR Creation',
                'User': bmr.created_by.get_full_name() if bmr.created_by else 'Unknown',
                'User Role': bmr.created_by.role if bmr.created_by else 'Unknown',
                'Date': bmr.created_date.strftime('%Y-%m-%d %H:%M:%S') if bmr.created_date else '',
                'Comments': bmr.qa_comments,
                'Status': bmr.status
            })
        
        if bmr.regulatory_comments:
            comments_data.append({
                'BMR Number': bmr.batch_number,
                'Product': bmr.product.product_name,
                'Comment Type': 'BMR Regulatory Comments',
                'Phase': 'Regulatory Approval',
                'User': bmr.approved_by.get_full_name() if bmr.approved_by else 'Unknown',
                'User Role': bmr.approved_by.role if bmr.approved_by else 'Unknown',
                'Date': (bmr.approved_date or bmr.modified_date).strftime('%Y-%m-%d %H:%M:%S') if (bmr.approved_date or bmr.modified_date) else '',
                'Comments': bmr.regulatory_comments,
                'Status': bmr.status
            })
    
    # Phase Level Comments
    if is_admin:
        phases = BatchPhaseExecution.objects.all().select_related(
            'bmr', 'bmr__product', 'phase', 'started_by', 'completed_by'
        )
    else:
        phases = BatchPhaseExecution.objects.filter(
            Q(started_by=request.user) | 
            Q(completed_by=request.user) |
            Q(bmr__created_by=request.user)
        ).select_related(
            'bmr', 'bmr__product', 'phase', 'started_by', 'completed_by'
        )
    
    for phase in phases:
        if phase.operator_comments:
            comments_data.append({
                'BMR Number': phase.bmr.batch_number,
                'Product': phase.bmr.product.product_name,
                'Comment Type': 'Operator Comments',
                'Phase': phase.phase.get_phase_name_display(),
                'User': phase.completed_by.get_full_name() if phase.completed_by else 
                        phase.started_by.get_full_name() if phase.started_by else 'Unknown',
                'User Role': (phase.completed_by.role if phase.completed_by else 
                            phase.started_by.role if phase.started_by else 'Unknown'),
                'Date': (phase.completed_date or phase.started_date or phase.created_date).strftime('%Y-%m-%d %H:%M:%S') if (phase.completed_date or phase.started_date or phase.created_date) else '',
                'Comments': phase.operator_comments,
                'Status': phase.status
            })
        
        if phase.qa_comments:
            comments_data.append({
                'BMR Number': phase.bmr.batch_number,
                'Product': phase.bmr.product.product_name,
                'Comment Type': 'Phase QA Comments',
                'Phase': phase.phase.get_phase_name_display(),
                'User': phase.completed_by.get_full_name() if phase.completed_by else 'Unknown',
                'User Role': phase.completed_by.role if phase.completed_by else 'Unknown',
                'Date': (phase.completed_date or phase.created_date).strftime('%Y-%m-%d %H:%M:%S') if (phase.completed_date or phase.created_date) else '',
                'Comments': phase.qa_comments,
                'Status': phase.status
            })
        
        if phase.rejection_reason:
            comments_data.append({
                'BMR Number': phase.bmr.batch_number,
                'Product': phase.bmr.product.product_name,
                'Comment Type': 'Rejection Reason',
                'Phase': phase.phase.get_phase_name_display(),
                'User': phase.completed_by.get_full_name() if phase.completed_by else 'Unknown',
                'User Role': phase.completed_by.role if phase.completed_by else 'Unknown',
                'Date': (phase.completed_date or phase.created_date).strftime('%Y-%m-%d %H:%M:%S') if (phase.completed_date or phase.created_date) else '',
                'Comments': phase.rejection_reason,
                'Status': phase.status
            })
    
    # Signature Comments
    if is_admin:
        signatures = BMRSignature.objects.all().select_related('bmr', 'bmr__product', 'signed_by')
    else:
        signatures = BMRSignature.objects.filter(
            Q(signed_by=request.user) |
            Q(bmr__created_by=request.user)
        ).select_related('bmr', 'bmr__product', 'signed_by')
    
    for signature in signatures:
        if signature.comments:
            comments_data.append({
                'BMR Number': signature.bmr.batch_number,
                'Product': signature.bmr.product.product_name,
                'Comment Type': 'Electronic Signature',
                'Phase': f"Signature - {signature.signed_by_role}",
                'User': signature.signed_by.get_full_name() if signature.signed_by else 'Unknown',
                'User Role': signature.signed_by_role,
                'Date': signature.signed_date.strftime('%Y-%m-%d %H:%M:%S') if signature.signed_date else '',
                'Comments': signature.comments,
                'Status': 'Signed'
            })
    
    # Create CSV response
    response = HttpResponse(content_type='text/csv')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    response['Content-Disposition'] = f'attachment; filename="KPI_Comments_Report_{timestamp}.csv"'
    
    if comments_data:
        writer = csv.DictWriter(response, fieldnames=comments_data[0].keys())
        writer.writeheader()
        writer.writerows(comments_data)
    else:
        writer = csv.writer(response)
        writer.writerow(['No comments found in the system'])
    
    return response

@login_required
def export_comments_word(request):
    """Export comments to Word format with role-based filtering"""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        return HttpResponse("python-docx library not installed. Please install it to use Word export.", 
                          content_type="text/plain")
    
    # Check if user is admin/staff - they see all comments
    is_admin = request.user.is_staff or request.user.is_superuser or request.user.role == 'admin'
    
    # Collect all comments (same logic as CSV)
    comments_data = []
    
    # BMR Level Comments
    if is_admin:
        bmrs = BMR.objects.all().select_related('product', 'created_by', 'approved_by')
    else:
        bmrs = BMR.objects.filter(
            Q(created_by=request.user) | 
            Q(approved_by=request.user)
        ).select_related('product', 'created_by', 'approved_by')
    
    # Collect comments (similar to CSV logic but simplified for Word)
    for bmr in bmrs:
        if bmr.qa_comments:
            comments_data.append({
                'bmr': bmr.batch_number,
                'product': bmr.product.product_name,
                'type': 'BMR QA Comments',
                'user': bmr.created_by.get_full_name() if bmr.created_by else 'Unknown',
                'date': bmr.created_date,
                'comments': bmr.qa_comments
            })
    
    # Create Word document
    doc = Document()
    doc.add_heading('KPI Comments Report', 0)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Total Comments: {len(comments_data)}')
    
    for comment in comments_data:
        doc.add_heading(f"BMR: {comment['bmr']} - {comment['product']}", level=1)
        doc.add_paragraph(f"Type: {comment['type']}")
        doc.add_paragraph(f"User: {comment['user']}")
        doc.add_paragraph(f"Date: {comment['date'].strftime('%Y-%m-%d %H:%M:%S') if comment['date'] else 'N/A'}")
        doc.add_paragraph(f"Comments: {comment['comments']}")
        doc.add_paragraph("")  # Add spacing
    
    # Prepare response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    response['Content-Disposition'] = f'attachment; filename="KPI_Comments_Report_{timestamp}.docx"'
    
    doc.save(response)
    return response

@login_required
def export_comments_excel(request):
    """Export comments to Excel format with role-based filtering"""
    try:
        import pandas as pd
        from io import BytesIO
    except ImportError:
        return HttpResponse("pandas library not installed. Please install it to use Excel export.", 
                          content_type="text/plain")
    
    # Check if user is admin/staff - they see all comments
    is_admin = request.user.is_staff or request.user.is_superuser or request.user.role == 'admin'
    
    # Collect all comments (same logic as CSV)
    comments_data = []
    
    # BMR Level Comments
    if is_admin:
        bmrs = BMR.objects.all().select_related('product', 'created_by', 'approved_by')
    else:
        bmrs = BMR.objects.filter(
            Q(created_by=request.user) | 
            Q(approved_by=request.user)
        ).select_related('product', 'created_by', 'approved_by')
    
    for bmr in bmrs:
        if bmr.qa_comments:
            comments_data.append({
                'BMR Number': bmr.batch_number,
                'Product': bmr.product.product_name,
                'Comment Type': 'BMR QA Comments',
                'Phase': 'BMR Creation',
                'User': bmr.created_by.get_full_name() if bmr.created_by else 'Unknown',
                'User Role': bmr.created_by.role if bmr.created_by else 'Unknown',
                'Date': bmr.created_date,
                'Comments': bmr.qa_comments,
                'Status': bmr.status
            })
    
    # Create Excel file
    output = BytesIO()
    if comments_data:
        df = pd.DataFrame(comments_data)
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Comments Report', index=False)
    else:
        # Create empty dataframe with headers
        df = pd.DataFrame(columns=['BMR Number', 'Product', 'Comment Type', 'Phase', 'User', 'User Role', 'Date', 'Comments', 'Status'])
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Comments Report', index=False)
    
    output.seek(0)
    
    # Prepare response
    response = HttpResponse(
        output.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    response['Content-Disposition'] = f'attachment; filename="KPI_Comments_Report_{timestamp}.xlsx"'
    
    return response

@login_required
def bmr_comments_detail(request, bmr_id):
    """Detailed view of all comments for a specific BMR with role-based access"""
    bmr = get_object_or_404(BMR, id=bmr_id)
    
    # Check if user has access to this BMR
    is_admin = request.user.is_staff or request.user.is_superuser or request.user.role == 'admin'
    
    if not is_admin:
        # Non-admin users can only view BMRs they created or were involved in
        user_bmrs = BMR.objects.filter(
            Q(created_by=request.user) | 
            Q(approved_by=request.user)
        )
        
        # Also check if user was involved in any phases of this BMR
        user_phases = BatchPhaseExecution.objects.filter(
            Q(started_by=request.user) | Q(completed_by=request.user),
            bmr=bmr
        )
        
        # Also check signatures
        user_signatures = BMRSignature.objects.filter(
            bmr=bmr,
            signed_by=request.user
        )
        
        # If user has no connection to this BMR, deny access
        if not (bmr in user_bmrs or user_phases.exists() or user_signatures.exists()):
            from django.contrib import messages
            messages.error(request, 'Access denied. You can only view BMRs you were involved in.')
            return redirect('reports:comments_report')
    
    # Collect all comments for this BMR
    comments = []
    
    # BMR level comments
    if bmr.qa_comments:
        comments.append({
            'type': 'BMR QA Comments',
            'phase': 'BMR Creation',
            'user': bmr.created_by.get_full_name() if bmr.created_by else 'Unknown',
            'role': bmr.created_by.role if bmr.created_by else 'Unknown',
            'date': bmr.created_date,
            'content': bmr.qa_comments,
            'status': bmr.status
        })
    
    if bmr.regulatory_comments:
        comments.append({
            'type': 'BMR Regulatory Comments',
            'phase': 'Regulatory Approval',
            'user': bmr.approved_by.get_full_name() if bmr.approved_by else 'Unknown',
            'role': bmr.approved_by.role if bmr.approved_by else 'Unknown',
            'date': bmr.approved_date or bmr.modified_date,
            'content': bmr.regulatory_comments,
            'status': bmr.status
        })
    
    # Phase comments
    phases = BatchPhaseExecution.objects.filter(bmr=bmr).select_related('phase', 'started_by', 'completed_by')
    
    for phase in phases:
        if phase.operator_comments:
            comments.append({
                'type': 'Operator Comments',
                'phase': phase.phase.get_phase_name_display(),
                'user': phase.completed_by.get_full_name() if phase.completed_by else 
                        phase.started_by.get_full_name() if phase.started_by else 'Unknown',
                'role': (phase.completed_by.role if phase.completed_by else 
                        phase.started_by.role if phase.started_by else 'Unknown'),
                'date': phase.completed_date or phase.started_date or phase.created_date,
                'content': phase.operator_comments,
                'status': phase.status
            })
        
        if phase.qa_comments:
            comments.append({
                'type': 'Phase QA Comments',
                'phase': phase.phase.get_phase_name_display(),
                'user': phase.completed_by.get_full_name() if phase.completed_by else 'Unknown',
                'role': phase.completed_by.role if phase.completed_by else 'Unknown',
                'date': phase.completed_date or phase.created_date,
                'content': phase.qa_comments,
                'status': phase.status
            })
        
        if phase.rejection_reason:
            comments.append({
                'type': 'Rejection Reason',
                'phase': phase.phase.get_phase_name_display(),
                'user': phase.completed_by.get_full_name() if phase.completed_by else 'Unknown',
                'role': phase.completed_by.role if phase.completed_by else 'Unknown',
                'date': phase.completed_date or phase.created_date,
                'content': phase.rejection_reason,
                'status': phase.status
            })
    
    # Signature comments
    signatures = BMRSignature.objects.filter(bmr=bmr).select_related('signed_by')
    
    for signature in signatures:
        if signature.comments:
            comments.append({
                'type': 'Electronic Signature',
                'phase': f"Signature - {signature.signed_by_role}",
                'user': signature.signed_by.get_full_name() if signature.signed_by else 'Unknown',
                'role': signature.signed_by_role,
                'date': signature.signed_date,
                'content': signature.comments,
                'status': 'Signed'
            })
    
    # Sort by date
    comments.sort(key=lambda x: x['date'] or datetime.min)
    
    context = {
        'bmr': bmr,
        'comments': comments,
        'total_comments': len(comments),
        'is_admin': is_admin,
        'user_role': request.user.role
    }
    
    return render(request, 'reports/bmr_comments_detail.html', context)
