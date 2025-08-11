#!/usr/bin/env python
"""
Comments Report Generator for Kampala Pharmaceutical Industries
Collects all comments from BMRs, phases, and quality control activities
"""
import os
import django
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Inches
from django.db.models import Q

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'kampala_pharma.settings')
django.setup()

from bmr.models import BMR, BMRSignature
from workflow.models import BatchPhaseExecution
from accounts.models import CustomUser

def collect_all_comments(user_id=None, user_role=None):
    """Collect all comments from the system with optional user filtering"""
    print("🔍 COLLECTING COMMENTS FROM PHARMACEUTICAL WORKFLOW")
    if user_id:
        print(f"📋 Filtering for User ID: {user_id} (Role: {user_role})")
    else:
        print("📋 Collecting ALL comments (Admin view)")
    print("=" * 60)
    
    comments_data = []
    
    # Determine if this is an admin user
    is_admin = False
    if user_id:
        try:
            user = CustomUser.objects.get(id=user_id)
            is_admin = user.is_staff or user.is_superuser or user.role == 'admin'
        except CustomUser.DoesNotExist:
            print(f"❌ User with ID {user_id} not found")
            return []
    else:
        is_admin = True  # If no user specified, assume admin view
    
    # 1. BMR Level Comments
    print("\n📋 BMR LEVEL COMMENTS:")
    if is_admin:
        bmrs = BMR.objects.all().select_related('product', 'created_by', 'approved_by')
    else:
        # Filter for specific user
        bmrs = BMR.objects.filter(
            Q(created_by_id=user_id) | Q(approved_by_id=user_id)
        ).select_related('product', 'created_by', 'approved_by')
    
    for bmr in bmrs:
        # QA Comments on BMR
        if bmr.qa_comments:
            comments_data.append({
                'bmr_number': bmr.batch_number,
                'product': bmr.product.product_name,
                'comment_type': 'BMR QA Comments',
                'phase': 'BMR Creation',
                'user': bmr.created_by.get_full_name() if bmr.created_by else 'Unknown',
                'user_role': bmr.created_by.role if bmr.created_by else 'Unknown',
                'date': bmr.created_date,
                'comments': bmr.qa_comments,
                'status': bmr.status
            })
            print(f"  ✅ QA Comments found for BMR {bmr.batch_number}")
        
        # Regulatory Comments on BMR
        if bmr.regulatory_comments:
            comments_data.append({
                'bmr_number': bmr.batch_number,
                'product': bmr.product.product_name,
                'comment_type': 'BMR Regulatory Comments',
                'phase': 'Regulatory Approval',
                'user': bmr.approved_by.get_full_name() if bmr.approved_by else 'Unknown',
                'user_role': bmr.approved_by.role if bmr.approved_by else 'Unknown',
                'date': bmr.approved_date or bmr.modified_date,
                'comments': bmr.regulatory_comments,
                'status': bmr.status
            })
            print(f"  ✅ Regulatory Comments found for BMR {bmr.batch_number}")
    
    # 2. Phase Level Comments
    print(f"\n⚙️ PHASE LEVEL COMMENTS:")
    if is_admin:
        phases = BatchPhaseExecution.objects.all().select_related(
            'bmr', 'bmr__product', 'phase', 'started_by', 'completed_by'
        )
    else:
        # Filter for specific user
        phases = BatchPhaseExecution.objects.filter(
            Q(started_by_id=user_id) | Q(completed_by_id=user_id) | Q(bmr__created_by_id=user_id)
        ).select_related(
            'bmr', 'bmr__product', 'phase', 'started_by', 'completed_by'
        )
    
    for phase in phases:
        # Operator Comments
        if phase.operator_comments:
            comments_data.append({
                'bmr_number': phase.bmr.batch_number,
                'product': phase.bmr.product.product_name,
                'comment_type': 'Operator Comments',
                'phase': phase.phase.get_phase_name_display(),
                'user': phase.completed_by.get_full_name() if phase.completed_by else 
                        phase.started_by.get_full_name() if phase.started_by else 'Unknown',
                'user_role': (phase.completed_by.role if phase.completed_by else 
                            phase.started_by.role if phase.started_by else 'Unknown'),
                'date': phase.completed_date or phase.started_date or phase.created_date,
                'comments': phase.operator_comments,
                'status': phase.status
            })
            print(f"  ✅ Operator Comments: {phase.bmr.batch_number} - {phase.phase.phase_name}")
        
        # QA Comments on Phase
        if phase.qa_comments:
            comments_data.append({
                'bmr_number': phase.bmr.batch_number,
                'product': phase.bmr.product.product_name,
                'comment_type': 'Phase QA Comments',
                'phase': phase.phase.get_phase_name_display(),
                'user': phase.completed_by.get_full_name() if phase.completed_by else 'Unknown',
                'user_role': phase.completed_by.role if phase.completed_by else 'Unknown',
                'date': phase.completed_date or phase.created_date,
                'comments': phase.qa_comments,
                'status': phase.status
            })
            print(f"  ✅ QA Comments: {phase.bmr.batch_number} - {phase.phase.phase_name}")
        
        # Rejection Reasons
        if phase.rejection_reason:
            comments_data.append({
                'bmr_number': phase.bmr.batch_number,
                'product': phase.bmr.product.product_name,
                'comment_type': 'Rejection Reason',
                'phase': phase.phase.get_phase_name_display(),
                'user': phase.completed_by.get_full_name() if phase.completed_by else 'Unknown',
                'user_role': phase.completed_by.role if phase.completed_by else 'Unknown',
                'date': phase.completed_date or phase.created_date,
                'comments': phase.rejection_reason,
                'status': phase.status
            })
            print(f"  ❌ Rejection Reason: {phase.bmr.batch_number} - {phase.phase.phase_name}")
    
    # 3. Signature Comments
    print(f"\n✍️ SIGNATURE COMMENTS:")
    if is_admin:
        signatures = BMRSignature.objects.all().select_related('bmr', 'bmr__product', 'signed_by')
    else:
        # Filter for specific user
        signatures = BMRSignature.objects.filter(
            Q(signed_by_id=user_id) | Q(bmr__created_by_id=user_id)
        ).select_related('bmr', 'bmr__product', 'signed_by')
    
    for signature in signatures:
        if signature.comments:
            comments_data.append({
                'bmr_number': signature.bmr.batch_number,
                'product': signature.bmr.product.product_name,
                'comment_type': 'Electronic Signature',
                'phase': f"Signature - {signature.signed_by_role}",
                'user': signature.signed_by.get_full_name() if signature.signed_by else 'Unknown',
                'user_role': signature.signed_by_role,
                'date': signature.signed_date,
                'comments': signature.comments,
                'status': 'Signed'
            })
            print(f"  ✍️ Signature Comments: {signature.bmr.batch_number} - {signature.signed_by_role}")
    
    print(f"\n📊 SUMMARY:")
    print(f"   • Total Comments Found: {len(comments_data)}")
    print(f"   • BMR Level: {len([c for c in comments_data if 'BMR' in c['comment_type']])}")
    print(f"   • Phase Level: {len([c for c in comments_data if c['comment_type'] in ['Operator Comments', 'Phase QA Comments', 'Rejection Reason']])}")
    print(f"   • Signatures: {len([c for c in comments_data if c['comment_type'] == 'Electronic Signature'])}")
    
    return comments_data

def export_to_excel(comments_data, filename=None):
    """Export comments to Excel file"""
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"KPI_Comments_Report_{timestamp}.xlsx"
    
    # Convert to DataFrame
    df = pd.DataFrame(comments_data)
    
    if df.empty:
        print("❌ No comments found to export")
        return None
    
    # Sort by BMR and date
    df = df.sort_values(['bmr_number', 'date'])
    
    # Create Excel writer with multiple sheets
    with pd.ExcelWriter(filename, engine='openpyxl') as writer:
        # All Comments Sheet
        df.to_excel(writer, sheet_name='All Comments', index=False)
        
        # BMR Comments Sheet
        bmr_comments = df[df['comment_type'].str.contains('BMR')]
        if not bmr_comments.empty:
            bmr_comments.to_excel(writer, sheet_name='BMR Comments', index=False)
        
        # Phase Comments Sheet
        phase_comments = df[df['comment_type'].isin(['Operator Comments', 'Phase QA Comments', 'Rejection Reason'])]
        if not phase_comments.empty:
            phase_comments.to_excel(writer, sheet_name='Phase Comments', index=False)
        
        # QC & Rejection Comments
        qc_comments = df[df['comment_type'].isin(['Rejection Reason', 'Phase QA Comments'])]
        if not qc_comments.empty:
            qc_comments.to_excel(writer, sheet_name='QC & Rejections', index=False)
        
        # By BMR Summary
        bmr_summary = df.groupby(['bmr_number', 'product']).agg({
            'comments': 'count',
            'comment_type': lambda x: ', '.join(x.unique())
        }).reset_index()
        bmr_summary.columns = ['BMR Number', 'Product', 'Total Comments', 'Comment Types']
        bmr_summary.to_excel(writer, sheet_name='BMR Summary', index=False)
    
    print(f"✅ Excel report exported: {filename}")
    return filename

def export_to_word(comments_data, filename=None):
    """Export comments to Word document"""
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"KPI_Comments_Report_{timestamp}.docx"
    
    # Create Word document
    doc = Document()
    
    # Title
    title = doc.add_heading('Kampala Pharmaceutical Industries', 0)
    title.alignment = 1  # Center alignment
    
    subtitle = doc.add_heading('Comments and Observations Report', level=1)
    subtitle.alignment = 1
    
    # Report metadata
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph(f"Total Comments: {len(comments_data)}")
    doc.add_paragraph(f"System: Pharmaceutical Operations Management")
    
    # Group comments by BMR
    bmr_groups = {}
    for comment in comments_data:
        bmr_num = comment['bmr_number']
        if bmr_num not in bmr_groups:
            bmr_groups[bmr_num] = []
        bmr_groups[bmr_num].append(comment)
    
    # Sort BMRs
    for bmr_num in sorted(bmr_groups.keys()):
        doc.add_page_break()
        
        # BMR Header
        bmr_heading = doc.add_heading(f'BMR: {bmr_num}', level=1)
        
        # BMR Info
        first_comment = bmr_groups[bmr_num][0]
        doc.add_paragraph(f"Product: {first_comment['product']}")
        doc.add_paragraph(f"Total Comments: {len(bmr_groups[bmr_num])}")
        
        # Comments for this BMR
        for i, comment in enumerate(sorted(bmr_groups[bmr_num], key=lambda x: x['date']), 1):
            # Comment header
            comment_heading = doc.add_heading(f'Comment {i}: {comment["comment_type"]}', level=2)
            
            # Comment details table
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Light Shading Accent 1'
            
            # Fill table
            cells = table.rows[0].cells
            cells[0].text = 'Phase:'
            cells[1].text = comment['phase']
            
            cells = table.rows[1].cells
            cells[0].text = 'User:'
            cells[1].text = f"{comment['user']} ({comment['user_role']})"
            
            cells = table.rows[2].cells
            cells[0].text = 'Date:'
            cells[1].text = comment['date'].strftime('%B %d, %Y at %I:%M %p') if comment['date'] else 'N/A'
            
            cells = table.rows[3].cells
            cells[0].text = 'Status:'
            cells[1].text = comment['status']
            
            cells = table.rows[4].cells
            cells[0].text = 'Type:'
            cells[1].text = comment['comment_type']
            
            cells = table.rows[5].cells
            cells[0].text = 'Comments:'
            cells[1].text = comment['comments'][:500] + ('...' if len(comment['comments']) > 500 else '')
            
            # Add full comments if truncated
            if len(comment['comments']) > 500:
                doc.add_heading('Full Comments:', level=3)
                doc.add_paragraph(comment['comments'])
            
            doc.add_paragraph()  # Add spacing
    
    # Save document
    doc.save(filename)
    print(f"✅ Word report exported: {filename}")
    return filename

def generate_summary_report():
    """Generate a summary of comments activity"""
    comments_data = collect_all_comments()
    
    if not comments_data:
        print("❌ No comments found in the system")
        return
    
    print(f"\n📈 DETAILED ANALYSIS:")
    
    # By Comment Type
    type_counts = {}
    for comment in comments_data:
        comment_type = comment['comment_type']
        type_counts[comment_type] = type_counts.get(comment_type, 0) + 1
    
    print(f"\n📝 Comments by Type:")
    for comment_type, count in sorted(type_counts.items()):
        print(f"   • {comment_type}: {count}")
    
    # By User Role
    role_counts = {}
    for comment in comments_data:
        role = comment['user_role']
        role_counts[role] = role_counts.get(role, 0) + 1
    
    print(f"\n👥 Comments by User Role:")
    for role, count in sorted(role_counts.items()):
        print(f"   • {role}: {count}")
    
    # By BMR
    bmr_counts = {}
    for comment in comments_data:
        bmr = comment['bmr_number']
        bmr_counts[bmr] = bmr_counts.get(bmr, 0) + 1
    
    print(f"\n📋 Comments by BMR (Top 10):")
    sorted_bmrs = sorted(bmr_counts.items(), key=lambda x: x[1], reverse=True)[:10]
    for bmr, count in sorted_bmrs:
        print(f"   • {bmr}: {count} comments")
    
    # Recent Activity
    recent_comments = sorted(comments_data, key=lambda x: x['date'] or datetime.min, reverse=True)[:10]
    print(f"\n🕒 Recent Comments (Last 10):")
    for comment in recent_comments:
        date_str = comment['date'].strftime('%Y-%m-%d %H:%M') if comment['date'] else 'N/A'
        print(f"   • {date_str} | {comment['bmr_number']} | {comment['user']} | {comment['comment_type']}")
    
    return comments_data

def main():
    """Main function to run the comments report"""
    print("🏭 KAMPALA PHARMACEUTICAL INDUSTRIES")
    print("📊 COMMENTS & OBSERVATIONS REPORT GENERATOR")
    print("=" * 60)
    
    # Ask if this is for a specific user or admin view
    print(f"\n👤 USER ACCESS LEVEL:")
    print("   1. Admin View (All comments)")
    print("   2. Operator View (Specific user's comments)")
    
    choice = input("\nSelect access level (1/2): ").strip()
    
    user_id = None
    user_role = None
    
    if choice == '2':
        # Get user list
        users = CustomUser.objects.filter(is_active=True).order_by('username')
        if not users:
            print("❌ No active users found in the system.")
            return
        
        print(f"\n👥 AVAILABLE USERS:")
        for i, user in enumerate(users, 1):
            print(f"   {i}. {user.get_full_name()} ({user.username}) - {user.get_role_display()}")
        
        try:
            user_choice = int(input(f"\nSelect user (1-{len(users)}): ").strip())
            if 1 <= user_choice <= len(users):
                selected_user = users[user_choice - 1]
                user_id = selected_user.id
                user_role = selected_user.role
                print(f"📋 Selected: {selected_user.get_full_name()} ({selected_user.get_role_display()})")
            else:
                print("❌ Invalid selection. Using admin view.")
        except ValueError:
            print("❌ Invalid input. Using admin view.")
    
    # Collect comments based on access level
    comments_data = collect_all_comments(user_id, user_role)
    
    if not comments_data:
        print(f"\n❌ No comments found.")
        if user_id:
            print("   This could mean:")
            print("   • This user hasn't made any comments yet")
            print("   • This user hasn't been involved in any BMRs or phases")
        else:
            print("   This could mean:")
            print("   • No BMRs have been created yet")
            print("   • No phases have been executed with comments")
            print("   • No quality control activities have been performed")
        return
    
    # Generate summary
    print(f"\n📈 DETAILED ANALYSIS:")
    
    # By Comment Type
    type_counts = {}
    for comment in comments_data:
        comment_type = comment['comment_type']
        type_counts[comment_type] = type_counts.get(comment_type, 0) + 1
    
    print(f"\n📝 Comments by Type:")
    for comment_type, count in sorted(type_counts.items()):
        print(f"   • {comment_type}: {count}")
    
    # By User Role
    role_counts = {}
    for comment in comments_data:
        role = comment['user_role']
        role_counts[role] = role_counts.get(role, 0) + 1
    
    print(f"\n👥 Comments by User Role:")
    for role, count in sorted(role_counts.items()):
        print(f"   • {role}: {count}")
    
    # By BMR
    bmr_counts = {}
    for comment in comments_data:
        bmr = comment['bmr_number']
        bmr_counts[bmr] = bmr_counts.get(bmr, 0) + 1
    
    print(f"\n📋 Comments by BMR (Top 10):")
    sorted_bmrs = sorted(bmr_counts.items(), key=lambda x: x[1], reverse=True)[:10]
    for bmr, count in sorted_bmrs:
        print(f"   • {bmr}: {count} comments")
    
    print(f"\n📋 EXPORT OPTIONS:")
    print("   1. Excel Report (Multiple sheets, data analysis)")
    print("   2. Word Report (Formatted document)")
    print("   3. Both Excel and Word")
    
    choice = input("\nSelect export option (1/2/3): ").strip()
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    if choice in ['1', '3']:
        excel_file = export_to_excel(comments_data, f"KPI_Comments_Report_{timestamp}.xlsx")
        if excel_file:
            print(f"📊 Excel file ready: {excel_file}")
    
    if choice in ['2', '3']:
        word_file = export_to_word(comments_data, f"KPI_Comments_Report_{timestamp}.docx")
        if word_file:
            print(f"📄 Word file ready: {word_file}")
    
    print(f"\n✅ REPORT GENERATION COMPLETE!")
    print(f"   • Total Comments Exported: {len(comments_data)}")
    print(f"   • Files saved in current directory")
    print(f"   • Timestamp: {timestamp}")

if __name__ == "__main__":
    main()
